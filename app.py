import streamlit as st
import scipy  # Needed for clustermap
import pandas as pd
import matplotlib.pyplot as plt
from data_cleaner import clean_data
from plot_generator import generate_plot
from insight_generator import init_gemini, generate_insight

from io import BytesIO
from docx import Document
from docx.shared import Inches

st.set_page_config(page_title="AutoVizAI", layout="wide")
st.title("📊 VizAI - Smart Analyzer with AI")

uploaded_file = st.file_uploader("📂 Upload your Excel/CSV file", type=["csv", "xlsx"])

if uploaded_file:
    file_ext = uploaded_file.name.split('.')[-1]

    try:
        if file_ext == 'csv':
            df = pd.read_csv(uploaded_file)
        else:
            df = pd.read_excel(uploaded_file)

        st.subheader("🔍 Raw Data Preview")
        st.write(df.head())

        st.subheader("🧼 Cleaning Data")
        cleaned_df = clean_data(df)
        st.success("✅ Data cleaned successfully!")

        st.subheader("📊 Plot Your Data")

        chart_types = [
            "scatterplot", "lineplot", "barplot", "countplot", "boxplot", "violinplot",
            "stripplot", "swarmplot", "histplot", "kdeplot",
            "pairplot", "heatmap", "clustermap", "jointplot", "rugplot"
        ]

        chart_type = st.selectbox("Choose a chart type", chart_types)
        x_axis = st.selectbox("X-axis", cleaned_df.columns) if chart_type not in ["pairplot", "clustermap", "heatmap"] else None
        y_axis = st.selectbox("Y-axis", cleaned_df.columns) if chart_type in [
            "scatterplot", "lineplot", "barplot", "boxplot", "violinplot",
            "stripplot", "swarmplot", "jointplot"
        ] else None

        if st.button("📈 Generate Plot"):
            fig = generate_plot(cleaned_df, chart_type, x_axis, y_axis)

            if fig and hasattr(fig, "savefig"):
                st.pyplot(fig)

                img_buf = BytesIO()
                fig.savefig(img_buf, format='png', bbox_inches='tight')
                img_buf.seek(0)

                try:
                    model = init_gemini()
                    insight = generate_insight(model, cleaned_df, chart_type, x_axis, y_axis, fig=fig)

                    st.subheader("🧠 AI Insight")
                    st.success("✅ Insight Generated:")
                    st.markdown(insight)

                    doc = Document()
                    doc.add_heading("AutoVizAI Report", level=1)
                    doc.add_paragraph(f"Chart Type: {chart_type}")
                    if x_axis:
                        doc.add_paragraph(f"X-axis: {x_axis}")
                    if y_axis:
                        doc.add_paragraph(f"Y-axis: {y_axis}")

                    doc.add_paragraph("AI Insight:")
                    doc.add_paragraph(insight)

                    doc.add_picture(img_buf, width=Inches(5.5))

                    output_buf = BytesIO()
                    doc.save(output_buf)
                    output_buf.seek(0)

                    st.download_button(
                        label="📥 Download Insight as DOCX",
                        data=output_buf,
                        file_name="AutoVizAI_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                except Exception as e:
                    st.error(f"❌ Failed to generate insight: {e}")
            else:
                st.warning("⚠️ Could not generate the plot. Please check your column selection.")
    except Exception as e:
        st.error(f"❌ Failed to process the file: {e}")
else:
    st.info("📥 Upload an Excel or CSV file to get started.")